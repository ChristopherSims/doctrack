"""
Document export functions for CSV, Word (.docx), and PDF formats.

Imports from database.py:
  - get_all_requirements
  - get_document_db
  - get_commits
  - get_traceability_links
"""

import csv
import io
import json
from datetime import datetime

from database import get_all_requirements, get_document_db, get_commits, get_traceability_links


def export_csv(document_id):
    """Export all requirements for a document as a CSV string.

    Columns: ID, Level, Title, Description, Status, Priority,
             Change Request ID, Change Request Link, Test Plan, Test Plan Link,
             Verification Method, Rationale, Tags, Created At, Updated At

    Args:
        document_id: The document ID to export.

    Returns:
        CSV string of all requirements, or an empty CSV with headers if none exist.
    """
    requirements = get_all_requirements(document_id)

    columns = [
        "ID", "Level", "Title", "Description", "Status", "Priority",
        "Change Request ID", "Change Request Link", "Test Plan", "Test Plan Link",
        "Verification Method", "Rationale", "Tags", "Created At", "Updated At"
    ]

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(columns)

    for req in requirements:
        # Tags are stored as a JSON string in the DB; convert to comma-separated
        tags_raw = req.get("tags", "[]")
        if isinstance(tags_raw, str):
            try:
                tags_list = json.loads(tags_raw)
            except (json.JSONDecodeError, TypeError):
                tags_list = []
        elif isinstance(tags_raw, list):
            tags_list = tags_raw
        else:
            tags_list = []
        tags_str = ", ".join(str(t) for t in tags_list)

        writer.writerow([
            req.get("id", ""),
            req.get("level", ""),
            req.get("title", ""),
            req.get("description", ""),
            req.get("status", ""),
            req.get("priority", ""),
            req.get("changeRequestId", ""),
            req.get("changeRequestLink", ""),
            req.get("testPlan", ""),
            req.get("testPlanLink", ""),
            req.get("verificationMethod", ""),
            req.get("rationale", ""),
            tags_str,
            req.get("createdAt", ""),
            req.get("updatedAt", ""),
        ])

    return output.getvalue()


def export_word(document_id):
    """Export a document and its requirements as a Word (.docx) file.

    Creates a title page with document metadata, a table-of-contents
    placeholder, a summary requirements table, and detailed sections
    per requirement.

    Args:
        document_id: The document ID to export.

    Returns:
        Bytes of the generated .docx file.

    Raises:
        ImportError: If python-docx is not installed.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise ImportError(
            "python-docx is required for Word export. "
            "Install it with: pip install python-docx>=0.8.11"
        )

    document_db = get_document_db(document_id)
    if not document_db:
        raise ValueError(f"Document with id '{document_id}' not found.")

    requirements = get_all_requirements(document_id)

    doc = Document()

    # ---- Title Page ----
    doc.add_paragraph("")  # spacer
    title_para = doc.add_heading(document_db.get("title", "Untitled Document"), level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_fields = [
        ("Version", document_db.get("version", "")),
        ("Status", document_db.get("status", "")),
        ("Owner", document_db.get("owner", "")),
        ("Created", document_db.get("createdAt", "")),
        ("Updated", document_db.get("updatedAt", "")),
    ]
    for label, value in meta_fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_value = p.add_run(str(value))
        run_value.font.size = Pt(11)

    doc.add_page_break()

    # ---- Table of Contents Placeholder ----
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph(
        "[Table of Contents — update this field in Word: right-click → Update Field]"
    )
    doc.add_page_break()

    # ---- Summary Requirements Table ----
    doc.add_heading("Requirements Summary", level=1)

    if requirements:
        summary_table = doc.add_table(rows=1, cols=4)
        summary_table.style = "Light Grid Accent 1"
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = "ID"
        hdr_cells[1].text = "Title"
        hdr_cells[2].text = "Status"
        hdr_cells[3].text = "Priority"

        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for req in requirements:
            row_cells = summary_table.add_row().cells
            row_cells[0].text = str(req.get("id", ""))
            row_cells[1].text = str(req.get("title", ""))
            row_cells[2].text = str(req.get("status", ""))
            row_cells[3].text = str(req.get("priority", ""))
    else:
        doc.add_paragraph("No requirements found for this document.")

    doc.add_page_break()

    # ---- Detailed Requirement Sections ----
    doc.add_heading("Detailed Requirements", level=1)

    if not requirements:
        doc.add_paragraph("No requirements to display.")
    else:
        for req in requirements:
            req_id = req.get("id", "unknown")
            doc.add_heading(f"Requirement {req_id}", level=2)

            detail_fields = [
                ("Level", req.get("level", "")),
                ("Title", req.get("title", "")),
                ("Description", req.get("description", "")),
                ("Status", req.get("status", "")),
                ("Priority", req.get("priority", "")),
                ("Change Request ID", req.get("changeRequestId", "")),
                ("Change Request Link", req.get("changeRequestLink", "")),
                ("Test Plan", req.get("testPlan", "")),
                ("Test Plan Link", req.get("testPlanLink", "")),
                ("Verification Method", req.get("verificationMethod", "")),
                ("Rationale", req.get("rationale", "")),
                ("Created At", req.get("createdAt", "")),
                ("Updated At", req.get("updatedAt", "")),
            ]

            # Tags need special handling (JSON list → string)
            tags_raw = req.get("tags", "[]")
            if isinstance(tags_raw, str):
                try:
                    tags_list = json.loads(tags_raw)
                except (json.JSONDecodeError, TypeError):
                    tags_list = []
            elif isinstance(tags_raw, list):
                tags_list = tags_raw
            else:
                tags_list = []
            detail_fields.append(("Tags", ", ".join(str(t) for t in tags_list)))

            # Traceability links
            trace_links = get_traceability_links(req_id)
            if trace_links:
                link_descriptions = []
                for link in trace_links:
                    direction = "→" if link.get("sourceRequirementId") == req_id else "←"
                    link_descriptions.append(
                        f"{direction} {link.get('targetRequirementId', '')} "
                        f"({link.get('linkType', '')}) "
                        f"[doc: {link.get('targetDocumentId', '')}]"
                    )
                detail_fields.append(("Traceability", "; ".join(link_descriptions)))
            else:
                detail_fields.append(("Traceability", "None"))

            detail_table = doc.add_table(rows=len(detail_fields), cols=2)
            detail_table.style = "Light Grid Accent 1"
            for i, (label, value) in enumerate(detail_fields):
                cell_label = detail_table.rows[i].cells[0]
                cell_value = detail_table.rows[i].cells[1]
                cell_label.text = label
                cell_value.text = str(value) if value else ""
                for paragraph in cell_label.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            doc.add_paragraph("")  # spacing between requirements

    # ---- Return bytes ----
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def export_pdf(document_id):
    """Export a document and its requirements as a PDF file.

    Creates a title page with document info and a requirements table.

    Args:
        document_id: The document ID to export.

    Returns:
        Bytes of the generated PDF file.

    Raises:
        ImportError: If reportlab is not installed.
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        )
    except ImportError:
        raise ImportError(
            "reportlab is required for PDF export. "
            "Install it with: pip install reportlab>=4.0.0"
        )

    document_db = get_document_db(document_id)
    if not document_db:
        raise ValueError(f"Document with id '{document_id}' not found.")

    requirements = get_all_requirements(document_id)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Title"],
        fontSize=24,
        spaceAfter=20,
        alignment=1,  # center
    )
    meta_style = ParagraphStyle(
        "DocMeta",
        parent=styles["Normal"],
        fontSize=12,
        spaceAfter=6,
        alignment=1,  # center
    )

    elements = []

    # ---- Title Page ----
    elements.append(Spacer(1, 2 * inch))
    elements.append(Paragraph(document_db.get("title", "Untitled Document"), title_style))
    elements.append(Spacer(1, 0.5 * inch))

    meta_items = [
        f"Version: {document_db.get('version', '')}",
        f"Status: {document_db.get('status', '')}",
        f"Owner: {document_db.get('owner', '')}",
        f"Created: {document_db.get('createdAt', '')}",
        f"Updated: {document_db.get('updatedAt', '')}",
    ]
    for item in meta_items:
        elements.append(Paragraph(item, meta_style))

    elements.append(PageBreak())

    # ---- Requirements Table ----
    heading_style = ParagraphStyle(
        "ReqHeading",
        parent=styles["Heading1"],
        fontSize=18,
        spaceAfter=12,
    )
    elements.append(Paragraph("Requirements", heading_style))
    elements.append(Spacer(1, 0.25 * inch))

    if not requirements:
        elements.append(Paragraph("No requirements found for this document.", styles["Normal"]))
    else:
        # Table header
        table_data = [["ID", "Level", "Title", "Description", "Status", "Priority"]]

        for req in requirements:
            # Truncate long descriptions for the table
            desc = req.get("description", "")
            if len(desc) > 100:
                desc = desc[:97] + "..."

            table_data.append([
                str(req.get("id", "")),
                str(req.get("level", "")),
                str(req.get("title", "")),
                str(desc),
                str(req.get("status", "")),
                str(req.get("priority", "")),
            ])

        # Column widths: ID, Level, Title, Description, Status, Priority
        col_widths = [0.8 * inch, 0.5 * inch, 1.2 * inch, 2.8 * inch, 0.7 * inch, 0.7 * inch]

        req_table = Table(table_data, colWidths=col_widths, repeatRows=1)
        req_table.setStyle(TableStyle([
            # Header styling
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2C3E50")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 10),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
            ("TOPPADDING", (0, 0), (-1, 0), 8),

            # Body styling
            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 1), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 1), (-1, -1), 4),
            ("TOPPADDING", (0, 1), (-1, -1), 4),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),

            # Alternating row colors
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#ECF0F1")]),

            # Grid
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#BDC3C7")),
        ]))

        elements.append(req_table)

    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer.read()
